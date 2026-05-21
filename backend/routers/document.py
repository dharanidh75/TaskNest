import os
import io
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from database import get_db, Folder, Note, Resource
from auth import get_current_user, User

router = APIRouter(tags=["document"])


def _build_summary_text(folder: Folder, notes: list, rag_summary: str) -> str:
    lines = [
        f"PROJECT: {folder.name}",
        f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC",
        "",
        "=" * 60,
        "PROJECT SUMMARY",
        "=" * 60,
        rag_summary,
        "",
    ]
    if notes:
        lines += ["=" * 60, "NOTES", "=" * 60]
        for n in notes:
            lines += [f"\n[ {n.title} ]", n.content or "", ""]
    return "\n".join(lines)


def _make_docx(folder: Folder, notes: list, rag_summary: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(folder.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Generated on {datetime.utcnow().strftime('%B %d, %Y')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x8F, 0xA3)

    doc.add_heading("Project Summary", level=1)
    doc.add_paragraph(rag_summary)

    if notes:
        doc.add_heading("Notes", level=1)
        for n in notes:
            doc.add_heading(n.title or "Untitled", level=2)
            doc.add_paragraph(n.content or "")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _make_pdf(folder: Folder, notes: list, rag_summary: str) -> io.BytesIO:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.colors import HexColor

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        accent = HexColor("#2b6777")

        title_style = ParagraphStyle("title", parent=styles["Title"],
                                     textColor=accent, fontSize=22, spaceAfter=6)
        h1_style = ParagraphStyle("h1", parent=styles["Heading1"],
                                  textColor=accent, fontSize=14, spaceAfter=4)
        h2_style = ParagraphStyle("h2", parent=styles["Heading2"],
                                  textColor=HexColor("#1a4f5c"), fontSize=12, spaceAfter=4)
        body_style = ParagraphStyle("body", parent=styles["Normal"],
                                    fontSize=10, leading=16, spaceAfter=8)

        story = [
            Paragraph(folder.name, title_style),
            Paragraph(f"Generated {datetime.utcnow().strftime('%B %d, %Y')}", styles["Normal"]),
            Spacer(1, 0.5*cm),
            Paragraph("Project Summary", h1_style),
            Spacer(1, 0.2*cm),
        ]

        for para in rag_summary.split("\n\n"):
            story.append(Paragraph(para.replace("\n", "<br/>"), body_style))

        if notes:
            story += [Spacer(1, 0.5*cm), Paragraph("Notes", h1_style)]
            for n in notes:
                story.append(Paragraph(n.title or "Untitled", h2_style))
                story.append(Paragraph((n.content or "").replace("\n", "<br/>"), body_style))
                story.append(Spacer(1, 0.3*cm))

        doc.build(story)
        buf.seek(0)
        return buf

    except ImportError:
        raise HTTPException(500, "reportlab not installed. Run: uv add reportlab")


@router.post("/folders/{folder_id}/generate-document/")
def generate_document(
    folder_id: int,
    fmt: str = Query("docx", regex="^(docx|pdf)$"),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    folder = db.query(Folder).filter(Folder.id == folder_id, Folder.user_id == user.id).first()
    if not folder:
        raise HTTPException(404, "Folder not found")

    notes = db.query(Note).filter(Note.folder_id == folder_id).all()

    # 1. Try to use an existing generated summary (most recent)
    summary_notes = [n for n in notes if n.title.startswith("Summary —")]
    if summary_notes:
        # Use the most recent summary that was already generated
        summary_notes.sort(key=lambda n: n.title, reverse=True)
        rag_summary = summary_notes[0].content
        print(f"[Document Gen] Using existing summary from {summary_notes[0].title}")
    else:
        # No existing summary — generate one using the same logic as the chatbot
        from langchain_groq import ChatGroq
        from langchain_core.messages import SystemMessage, HumanMessage

        llm = ChatGroq(
            groq_api_key=os.getenv("GROQ_API_KEY"),
            model_name="llama-3.1-8b-instant",
            temperature=0.2,
        )

        resources = db.query(Resource).filter(Resource.folder_id == folder_id).all()

        # Retrieve ALL chunks from ChromaDB
        from rag.chroma_store import get_collection
        try:
            collection = get_collection(folder_id)
            chroma_data = collection.get()
            all_chunks = chroma_data.get("documents", []) if chroma_data else []
            print(f"[Document Gen] Retrieved {len(all_chunks)} chunks for folder {folder_id}")
        except Exception as e:
            print(f"[Document Gen] ChromaDB retrieval failed: {e}")
            all_chunks = []

        file_list = ", ".join([r.filename for r in resources]) if resources else "(no files)"
        indexed_count = sum(1 for r in resources if r.indexed) if resources else 0
        context = "\n\n".join(all_chunks[:20]) if all_chunks else "[No indexed resource content available]"

        notes_text = "\n".join([f"- {n.title}: {(n.content or '')[:300]}" for n in notes if not n.title.startswith("Summary —")])

        prompt = (
            f"You are a professional technical writer creating comprehensive project documentation for ResHub.\n"
            f"Project: '{folder.name}'\n"
            f"Files uploaded: {file_list}\n"
            f"Indexed: {indexed_count}/{len(resources)}\n\n"
            f"RESOURCE CONTENT FROM UPLOADED FILES:\n"
            f"--- START RESOURCES ---\n{context}\n--- END RESOURCES ---\n\n"
            f"PROJECT NOTES:\n{notes_text or '(None)'}\n\n"
            "TASK: Write a detailed, structured project summary covering:\n"
            "• Purpose and objectives\n"
            "• Key features and functionality\n"
            "• Technical details and architecture\n"
            "• Current status and progress\n"
            "• Important findings from resources\n\n"
            "STRICT RULES:\n"
            "- Only use information from the provided resource blocks.\n"
            "- If resources are missing, explicitly state that.\n"
            "- Do NOT generate default content about other projects.\n"
            "- Be thorough and professional."
        )

        response = llm.invoke([
            SystemMessage(content="You are a ResHub professional technical writer. Only synthesize from provided resources. Do not generate fabricated content."),
            HumanMessage(content=prompt),
        ])
        rag_summary = response.content

    safe_name = "".join(c if c.isalnum() or c in ("-", "_") else "_" for c in folder.name)

    if fmt == "pdf":
        buf = _make_pdf(folder, notes, rag_summary)
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}_document.pdf"'},
        )
    else:
        buf = _make_docx(folder, notes, rag_summary)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}_document.docx"'},
        )